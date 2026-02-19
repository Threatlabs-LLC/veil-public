"""Document text extraction — PDF, DOCX, TXT, CSV, XLSX.

Extracts text from uploaded files for sanitization. Files are processed
in-memory and never stored — only metadata is persisted.
"""

from __future__ import annotations

from dataclasses import dataclass

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_TEXT_CHARS = 500_000  # 500K character limit
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx"}


@dataclass
class ExtractedDocument:
    """Result of extracting text from a document."""

    text: str
    filename: str
    file_type: str
    file_size_bytes: int
    char_count: int
    page_count: int | None = None
    was_truncated: bool = False


class ExtractionError(Exception):
    """Raised when text extraction fails."""


def extract_text(data: bytes, filename: str) -> ExtractedDocument:
    """Extract text from a file. Dispatches to the correct extractor.

    Args:
        data: Raw file bytes.
        filename: Original filename (used to determine type).

    Returns:
        ExtractedDocument with the extracted text and metadata.

    Raises:
        ExtractionError: If the file type is unsupported or extraction fails.
        ValueError: If the file exceeds size limits.
    """
    if len(data) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File size {len(data)} bytes exceeds limit of {MAX_FILE_SIZE_BYTES // (1024 * 1024)} MB"
        )

    ext = _get_extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise ExtractionError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".txt": _extract_txt,
        ".csv": _extract_csv,
        ".xlsx": _extract_xlsx,
    }

    text, page_count = extractors[ext](data)

    was_truncated = False
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS]
        was_truncated = True

    return ExtractedDocument(
        text=text,
        filename=filename,
        file_type=ext.lstrip("."),
        file_size_bytes=len(data),
        char_count=len(text),
        page_count=page_count,
        was_truncated=was_truncated,
    )


def _get_extension(filename: str) -> str:
    """Get the lowercase file extension."""
    dot_idx = filename.rfind(".")
    if dot_idx == -1:
        return ""
    return filename[dot_idx:].lower()


def _extract_pdf(data: bytes) -> tuple[str, int | None]:
    """Extract text from a PDF file."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ExtractionError("pypdf is not installed. Install with: pip install pypdf")

    import io

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)
        return "\n\n".join(pages), len(reader.pages)
    except Exception as e:
        raise ExtractionError(f"Failed to extract text from PDF: {e}")


def _extract_docx(data: bytes) -> tuple[str, int | None]:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise ExtractionError("python-docx is not installed. Install with: pip install python-docx")

    import io

    try:
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs), None
    except Exception as e:
        raise ExtractionError(f"Failed to extract text from DOCX: {e}")


def _extract_txt(data: bytes) -> tuple[str, int | None]:
    """Extract text from a plain text file."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding), None
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode text file — unsupported encoding")


def _extract_csv(data: bytes) -> tuple[str, int | None]:
    """Extract text from a CSV file (returns as plain text)."""
    import csv
    import io

    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = data.decode("latin-1")
        except UnicodeDecodeError:
            raise ExtractionError("Could not decode CSV file")

    try:
        reader = csv.reader(io.StringIO(text))
        rows = []
        for row in reader:
            rows.append(", ".join(row))
        return "\n".join(rows), None
    except csv.Error as e:
        raise ExtractionError(f"Failed to parse CSV: {e}")


def _extract_xlsx(data: bytes) -> tuple[str, int | None]:
    """Extract text from an XLSX file."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ExtractionError(
            "openpyxl is not installed. Install with: pip install openpyxl"
        )

    import io

    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows.append(", ".join(cells))
            if rows:
                sheets.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(sheets), len(wb.sheetnames) if sheets else 0
    except Exception as e:
        raise ExtractionError(f"Failed to extract text from XLSX: {e}")
