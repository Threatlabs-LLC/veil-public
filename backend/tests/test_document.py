"""Tests for the document text extraction module."""

from __future__ import annotations

import csv
import io

import pytest

from backend.core.document import (
    MAX_FILE_SIZE_BYTES,
    ExtractionError,
    ExtractedDocument,
    extract_text,
)


# ---------------------------------------------------------------------------
# TXT extraction
# ---------------------------------------------------------------------------


class TestTxtExtraction:
    def test_plain_text(self):
        data = b"Hello world, this is a test document."
        result = extract_text(data, "test.txt")
        assert result.text == "Hello world, this is a test document."
        assert result.file_type == "txt"
        assert result.char_count == len("Hello world, this is a test document.")
        assert result.was_truncated is False

    def test_utf8_with_bom(self):
        data = b"\xef\xbb\xbfHello with BOM"
        result = extract_text(data, "bom.txt")
        assert "Hello with BOM" in result.text

    def test_latin1_encoding(self):
        data = "Caf\xe9 r\xe9sum\xe9".encode("latin-1")
        result = extract_text(data, "latin.txt")
        assert "Caf" in result.text

    def test_empty_text(self):
        result = extract_text(b"", "empty.txt")
        assert result.text == ""
        assert result.char_count == 0

    def test_multiline(self):
        data = b"Line 1\nLine 2\nLine 3"
        result = extract_text(data, "multi.txt")
        assert result.text == "Line 1\nLine 2\nLine 3"


# ---------------------------------------------------------------------------
# CSV extraction
# ---------------------------------------------------------------------------


class TestCsvExtraction:
    def test_basic_csv(self):
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["name", "email", "phone"])
        writer.writerow(["Alice", "alice@example.com", "555-1234"])
        writer.writerow(["Bob", "bob@example.com", "555-5678"])
        data = buf.getvalue().encode("utf-8")

        result = extract_text(data, "contacts.csv")
        assert result.file_type == "csv"
        assert "alice@example.com" in result.text
        assert "Bob" in result.text

    def test_csv_preserves_all_rows(self):
        rows = "a,b\n1,2\n3,4\n"
        result = extract_text(rows.encode(), "data.csv")
        assert "1, 2" in result.text
        assert "3, 4" in result.text


# ---------------------------------------------------------------------------
# File limits
# ---------------------------------------------------------------------------


class TestFileLimits:
    def test_oversized_file_rejected(self):
        data = b"x" * (MAX_FILE_SIZE_BYTES + 1)
        with pytest.raises(ValueError, match="exceeds limit"):
            extract_text(data, "huge.txt")

    def test_max_size_file_accepted(self):
        data = b"x" * MAX_FILE_SIZE_BYTES
        result = extract_text(data, "maxsize.txt")
        assert result.file_size_bytes == MAX_FILE_SIZE_BYTES

    def test_text_truncation(self):
        # Create text just over the char limit
        data = ("a" * 600_000).encode("utf-8")
        result = extract_text(data, "big.txt")
        assert result.was_truncated is True
        assert result.char_count == 500_000


# ---------------------------------------------------------------------------
# Unsupported types
# ---------------------------------------------------------------------------


class TestUnsupportedTypes:
    def test_unsupported_extension(self):
        with pytest.raises(ExtractionError, match="Unsupported file type"):
            extract_text(b"data", "file.zip")

    def test_no_extension(self):
        with pytest.raises(ExtractionError, match="Unsupported file type"):
            extract_text(b"data", "noext")

    def test_unknown_extension(self):
        with pytest.raises(ExtractionError, match="Unsupported file type"):
            extract_text(b"data", "file.pptx")


# ---------------------------------------------------------------------------
# PDF extraction (requires pypdf)
# ---------------------------------------------------------------------------


class TestPdfExtraction:
    @pytest.fixture(autouse=True)
    def check_pypdf(self):
        try:
            import pypdf  # noqa: F401
        except ImportError:
            pytest.skip("pypdf not installed")

    def test_minimal_pdf(self):
        """Create a minimal valid PDF and extract text."""
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        # pypdf can't easily add text to blank pages, so we test with a real PDF
        buf = io.BytesIO()
        writer.write(buf)
        data = buf.getvalue()

        result = extract_text(data, "blank.pdf")
        assert result.file_type == "pdf"
        assert result.page_count == 1

    def test_corrupted_pdf(self):
        with pytest.raises(ExtractionError, match="Failed to extract"):
            extract_text(b"not a real pdf", "bad.pdf")


# ---------------------------------------------------------------------------
# DOCX extraction (requires python-docx)
# ---------------------------------------------------------------------------


class TestDocxExtraction:
    @pytest.fixture(autouse=True)
    def check_docx(self):
        try:
            import docx  # noqa: F401
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_basic_docx(self):
        """Create a minimal DOCX and extract text."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph with test data.")
        doc.add_paragraph("Second paragraph with more content.")
        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()

        result = extract_text(data, "test.docx")
        assert result.file_type == "docx"
        assert "First paragraph" in result.text
        assert "Second paragraph" in result.text

    def test_corrupted_docx(self):
        with pytest.raises(ExtractionError, match="Failed to extract"):
            extract_text(b"not a real docx", "bad.docx")


# ---------------------------------------------------------------------------
# XLSX extraction (requires openpyxl)
# ---------------------------------------------------------------------------


class TestXlsxExtraction:
    @pytest.fixture(autouse=True)
    def check_openpyxl(self):
        try:
            import openpyxl  # noqa: F401
        except ImportError:
            pytest.skip("openpyxl not installed")

    def test_basic_xlsx(self):
        """Create a minimal XLSX and extract text."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Name", "Email"])
        ws.append(["Alice", "alice@test.com"])
        ws.append(["Bob", "bob@test.com"])
        buf = io.BytesIO()
        wb.save(buf)
        data = buf.getvalue()

        result = extract_text(data, "data.xlsx")
        assert result.file_type == "xlsx"
        assert "alice@test.com" in result.text
        assert "[Sheet: Data]" in result.text

    def test_corrupted_xlsx(self):
        with pytest.raises(ExtractionError, match="Failed to extract"):
            extract_text(b"not a real xlsx", "bad.xlsx")


# ---------------------------------------------------------------------------
# ExtractedDocument properties
# ---------------------------------------------------------------------------


class TestExtractedDocument:
    def test_fields(self):
        doc = ExtractedDocument(
            text="hello",
            filename="test.txt",
            file_type="txt",
            file_size_bytes=5,
            char_count=5,
        )
        assert doc.text == "hello"
        assert doc.filename == "test.txt"
        assert doc.page_count is None
        assert doc.was_truncated is False
